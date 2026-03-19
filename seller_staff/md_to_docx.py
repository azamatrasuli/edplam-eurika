#!/usr/bin/env python3
"""
Convert 8 markdown knowledge base files into a single DOCX document.

Reads: 01_company.md through 08_contacts.md
Outputs: "База знаний для ИИ агента продавца.docx"
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

KB_DIR = os.path.dirname(os.path.abspath(__file__))
KB_PATH = os.path.join(KB_DIR, "knowledge_base")
OUTPUT_PATH = os.path.join(KB_DIR, "База знаний для ИИ агента продавца.docx")

MD_FILES = [
    "01_company.md",
    "02_products.md",
    "03_attestation.md",
    "04_sales_scripts.md",
    "05_enrollment.md",
    "06_faq.md",
    "07_education.md",
    "08_contacts.md",
]


def add_formatted_runs(paragraph, text):
    """Parse inline markdown formatting (**bold**) and add runs to paragraph."""
    # Split by **bold** markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def add_table_to_doc(doc, rows_data):
    """Add a formatted table with borders to the document."""
    if not rows_data or len(rows_data) < 2:
        return

    header = rows_data[0]
    data_rows = rows_data[2:]  # skip separator row (index 1)

    num_cols = len(header)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    # Header row
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text.strip())
        run.bold = True
        run.font.size = Pt(9)
        # Gray background for header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        set_cell_border(cell)

    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        for col_idx in range(min(len(row_data), num_cols)):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            cell_text = row_data[col_idx].strip()
            add_formatted_runs(p, cell_text)
            for run in p.runs:
                run.font.size = Pt(9)
            set_cell_border(cell)

    doc.add_paragraph()  # spacing after table


def parse_table_row(line):
    """Parse a markdown table row into cells."""
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [cell.strip() for cell in line.split('|')]


def is_separator_row(line):
    """Check if a markdown table row is a separator (|---|---|)."""
    stripped = line.strip().replace('|', '').replace('-', '').replace(':', '').strip()
    return len(stripped) == 0 and '---' in line


def main():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # -- Default font --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # -- Title --
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('База знаний ИИ-агента "Эврика" — EdPalm / ЦПСО')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

    # -- Subtitle lines --
    subtitle_lines = [
        "Источник данных: hss.center",
        "Дата сборки: март 2026",
        "Цены: 2026-2027 учебный год",
    ]
    for line in subtitle_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # -- Sequential heading counter --
    h2_counter = 0

    # -- Process each MD file --
    for md_file in MD_FILES:
        filepath = os.path.join(KB_PATH, md_file)
        with open(filepath, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.rstrip('\n')

            # Skip empty lines
            if stripped.strip() == '':
                i += 1
                continue

            # Skip H1 (file title like "# База знаний: О компании")
            if stripped.startswith('# ') and not stripped.startswith('## '):
                i += 1
                continue

            # H2 heading -> numbered Heading 2
            if stripped.startswith('## ') and not stripped.startswith('### '):
                h2_counter += 1
                heading_text = stripped[3:].strip()
                doc.add_heading(f"{h2_counter}. {heading_text}", level=2)
                i += 1
                continue

            # H3 heading -> Heading 3
            if stripped.startswith('### '):
                heading_text = stripped[4:].strip()
                doc.add_heading(heading_text, level=3)
                i += 1
                continue

            # Table detection
            if '|' in stripped and stripped.strip().startswith('|'):
                table_rows = []
                while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                    row_line = lines[i].rstrip('\n')
                    if is_separator_row(row_line):
                        table_rows.append('SEPARATOR')
                    else:
                        table_rows.append(parse_table_row(row_line))
                    i += 1

                # Filter: separate header, separator, and data
                parsed_rows = []
                for r in table_rows:
                    if r == 'SEPARATOR':
                        parsed_rows.append('SEPARATOR')
                    else:
                        parsed_rows.append(r)

                if len(parsed_rows) >= 2:
                    add_table_to_doc(doc, parsed_rows)
                continue

            # Bullet list (- item)
            if re.match(r'^- ', stripped.strip()):
                text = stripped.strip()[2:]
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_runs(p, text)
                i += 1
                # Check for sub-bullets (  - item)
                while i < len(lines):
                    next_line = lines[i].rstrip('\n')
                    if re.match(r'^  - ', next_line):
                        sub_text = next_line.strip()[2:]
                        sp = doc.add_paragraph(style='List Bullet 2')
                        add_formatted_runs(sp, sub_text)
                        i += 1
                    elif re.match(r'^- ', next_line.strip()):
                        # Next top-level bullet - let outer loop handle it
                        break
                    else:
                        break
                continue

            # Numbered list (1. item, 2. item, etc.)
            num_match = re.match(r'^(\d+)\.\s+(.+)', stripped.strip())
            if num_match:
                text = num_match.group(2)
                p = doc.add_paragraph(style='List Number')
                add_formatted_runs(p, text)
                i += 1
                continue

            # Normal paragraph
            text = stripped.strip()
            if text:
                p = doc.add_paragraph()
                add_formatted_runs(p, text)
            i += 1

    # -- Save --
    doc.save(OUTPUT_PATH)
    print(f"DOCX saved: {OUTPUT_PATH}")
    print(f"Total H2 sections: {h2_counter}")


if __name__ == '__main__':
    main()
